from docx.shared import Pt, RGBColor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import asyncio
import nest_asyncio
from tqdm.notebook import tqdm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re


def create_multilevel_numbering(document):
    numbering = document.part.numbering_part.numbering_definitions._numbering
    numId = len(numbering.xpath('./w:num')) + 1

    abstractNumId = numId
    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(abstractNumId))

    for level in range(0, 9):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'decimal')
        lvl.append(numFmt)

        lvlText = OxmlElement('w:lvlText')
        lvl_numbers = '.'.join(['%{}'.format(i+1) for i in range(level+1)]) + '.'  # Добавляем точку в конец
        lvlText.set(qn('w:val'), lvl_numbers)
        lvl.append(lvlText)

        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        if level == 0:
            ind.set(qn('w:center'), str(360 * level))
        else:
            ind.set(qn('w:left'), str(360 * level))  # Настройка отступа для каждого уровня

        pPr.append(ind)
        lvl.append(pPr)

        # Добавляем настройки шрифта для нумерации
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.append(rFonts)
        lvl.append(rPr)

        abstractNum.append(lvl)

    numbering.append(abstractNum)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(numId))
    abstractNumId_el = OxmlElement('w:abstractNumId')
    abstractNumId_el.set(qn('w:val'), str(abstractNumId))
    num.append(abstractNumId_el)
    numbering.append(num)

    return numId

terms = ["ИНН", "ОГРН",'Юридический/почтовый адрес', "Банковские реквизиты", "Телефон", "Электронная почта"]
meta_data_law = {"ИП": "Индивидуальный предприниматель",
                 "ООО": "Общество с ограниченной ответственностью",
                 "АО": "Акционерное общество"}

DEFAULT_SYSTEM_PROMPT_2 = ''

def add_paragraph_with_numbering(document, text, numId, level, align='left'):
    p = document.add_paragraph()
    p_paragraph = p._p

    numPr = OxmlElement('w:numPr')

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)

    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(numId))
    numPr.append(numId_el)

    p_pPr = p_paragraph.get_or_add_pPr()
    p_pPr.append(numPr)

    p.add_run(text)
    
    if align == 'left':     p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)


def get_my(dict1, text):
    answer = dict1.get(text, "Неизвестно")
    return answer #if answer != "" else "Неизвестно"


def get_head_uslugi(obj, side):

    head = []
    head.append(meta_data_law.get(get_my(obj, f"Юридическая форма {side}"), get_my(obj, f"Юридическая форма {side}")))

    company = obj.get(f"Название организации {side}")
    if company:
        head.append(company)

    inn = obj.get(f"ИНН {side}")
    ogrn = obj.get(f"ОГРН {side}")
    egrip = obj.get(f"ЕГРИП {side}")
    if inn:
        head.append("ИНН " + inn )

    if ogrn:
        head.append("ОГРН " + ogrn)
    if egrip:
        head.append("ЕГРИП " + egrip)

    if obj.get(f"Юридическая форма {side}") in ["ООО", "АО"]:
      address = obj.get(f"Юридический/почтовый адрес {side}")
      if address:
        head.append("Юридический адрес:")
        head.append(address)
      fio = obj.get(f"ФИО {side}")
      if fio:
          head.append("в лице генерального директора")
          head.append(fio.title())
          head.append("действующего на основании устава")
    else:
      fio = obj.get(f"ФИО {side}")
      if fio:
          head.append(fio.title())
    return " ".join(head)

def get_head_dkp(json_result, side):
    # ИП - ФИО - Данные
    # ООО - Данные - Юр. Адрес - ФИО
    head = []
    head.append(meta_data_law.get(get_my(json_result, f"Юридическая форма {side}"), get_my(json_result, f"Юридическая форма {side}")))

    inn = json_result.get(f"ИНН {side}")
    ogrn = json_result.get(f"ОГРН {side}")
    egrip = json_result.get(f"ЕГРИП {side}")
    
    if json_result.get(f"Юридическая форма {side}") in ["ООО", "АО"]:
        company = json_result.get(f"Название организации {side}").replace('ООО', '').replace('АО', '').replace(
            'Общество с ограниченной отвественностью', '').replace('Акционерное общество', '')
        if company:
            head.append(company)
        if inn:
            head.append("ИНН " + inn )
        if ogrn:
            head.append("ОГРН " + ogrn)
        if egrip:
            head.append("ЕГРИП " + egrip)
            
        address = json_result.get(f"Юридический/почтовый адрес {side}")
        if address:
            head.append("Юридический адрес:")
            head.append(address)
        fio = json_result.get(f"ФИО {side}")
        if fio:
            head.append("в лице генерального директора")
            head.append(fio.title())
    else:
        fio = json_result.get(f"ФИО {side}")
        if fio:
            head.append(fio.title())
        if inn:
            head.append("ИНН " + inn )
        if ogrn:
            head.append("ОГРН " + ogrn)
        if egrip:
            head.append("ЕГРИП " + egrip)
   
    return " ".join(head)


def add_heading_numbered_dkp(doc, num, text, align='left'):
  text = text.upper()
  if num is None:
        subject = doc.add_heading(f"{text}", level=1)
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subject.runs[0]
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = False
  else:
        add_paragraph_with_numbering(doc, text, num, 0, align=align)


def add_list(doc, num, list_text):
    '''
    new_num = 1
    for text in list_text:
        subject = doc.add_paragraph(f"{num}.{new_num}. {text}",
                     style='List')
        subject.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = subject.runs[0]
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        new_num += 1
    '''
    for text in list_text:
        if text is not None:
            add_paragraph_with_numbering(doc, text, num, 1, align='justify')
        

def add_text(doc, text, rephrase=False):
    subject = doc.add_paragraph(text)
    subject.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = subject.runs[0]
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    

async def async_generate(model, tokenizer, texts, generation_config=None):
    resp = []
    key, sentences = texts[0], texts[1]
    for text in sentences:
        if text is None: continue
        print(text, end='\n\n')
        data = tokenizer.apply_chat_template([{"role": "user", "content": DEFAULT_SYSTEM_PROMPT_2 + text}], tokenize=False, add_generation_prompt=True)
        data = tokenizer(data, return_tensors="pt", add_special_tokens=False)
        data = {k: v.to(model.device) for k, v in data.items()}
        output_ids   = model.generate(**data, generation_config=generation_config)[0] if generation_config else model.generate(**data)[0]
        output_ids   = output_ids[len(data["input_ids"][0]):]
        result       = tokenizer.decode(output_ids, skip_special_tokens=True).split('Исправленный текст:')[-1].split('markdown')[-1].strip() 
        cleaned_text = re.sub(r'[^а-яА-Яa-zA-Z0-9\.\:\s,\-\«\»%{}\(\)]', '', result).strip() 
        resp.append(cleaned_text)
        print(cleaned_text)
        print('-------')
    return key, resp


async def generate_concurrently(llm, tokenizer, prompts, gen_config=None):
    tasks  = [async_generate(llm, tokenizer, prompt, generation_config=gen_config) for prompt in tqdm(prompts.items())]
    result =  await asyncio.gather(*tasks)
    return dict(result)

        
async def json_to_doc_dcp(json_result, path_to_file, model, tokenizer, gen_config=None):
    global DEFAULT_SYSTEM_PROMPT_2 
    DEFAULT_SYSTEM_PROMPT_2 = """
    Ты - юридический эксперт. 
    Твоя задача - исправлять ошибки в юридическом тексте, 
    вернув в исправленном виде и не потеряв ключевые юридические сущности. 
    Не возвращай никакого другого текста, кроме отредактированного текста! 
    С большой буквы всегда: "Договор", "Цена", "Товар", "Продавец", "Покупатель", "Сторона".
    Никаких комментариев, только отредактированный текст!
    Текст: 

    """
    
    def fast_select(key, entity, text_rep, text_temp, dicts=True):
        if get_my(json_result, entity) not in ['', "Неизвестно", 'неизвестно', 'нет данных']:
            if dicts: sentences_rephrase[key] = [text_rep]
            else: return text_rep
        else:
            if dicts: sentences_template[key] = [text_temp]
            else: return text_temp
    
    doc = Document()
    head_number = create_multilevel_numbering(doc)
    today = date.today()

    data = today.strftime("%d.%m.%Y")
    #head_number = 1
    add_heading_numbered_dkp(doc, None, 'ДОГОВОР КУПЛИ-ПРОДАЖИ ТОВАРА')
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'г. Москва'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = '\t' + data
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    
    sentences_template, sentences_rephrase = {}, {}
   
    sentences_rephrase['шапка'] = [f"""{get_head_dkp(json_result, "продавца")} именуемый в дальнейшем «Продавец»,
                                   с одной стороны,\n {get_head_dkp(json_result, "покупателя")} именуемый в дальнейшем «Покупатель», 
                                   с другой стороны',\n именуемые вместе «Стороны», заключили настоящий договор (далее «Договор») о нижеследующем:\n\n"""]
    
    sentences_rephrase['предмет договора'] = [f"Продавец обязуется передать Покупателю {get_my(json_result, 'Товар')} {get_my(json_result, 'Информация о товаре')}  (именуемый в дальнейшем «Товар»)."]
    
    where = fast_select('адрес', 'Адрес доставки',
                f"по адресу {get_my(json_result, 'Адрес доставки')}",
                "по адресу фактического местонахождения Покупателя", dicts=False)
    
    when = fast_select('когда', 'Срок',
                f"в срок адресу {get_my(json_result, 'Срок')}",
                "в разумный срок", dicts=False)   
    
    sentences_rephrase['доставка'] = [f"Продавец обязуется обеспечить доставку товара по адресу {where} в срок {when}."]
    sentences_rephrase['цена']     = [f"Покупатель обязуется уплатить Продавцу Цену за приобретение Товара в размере {get_my(json_result, 'Цена товара')} {get_my(json_result, 'Валюта товара')} (далее – «Цена») в порядке, установленном настоящей статьей."]
                                   
    pay_method  = get_my(json_result, 'Способ оплаты товара').replace('неизвестно', 'безналичным').replace('Неизвестно', 'безналичным')
    if 'безнал' in pay_method:
        sentences_template['pay_phrase']  = "Уплата цены Договора производится путем безналичного перечисления денежных средств на счет Продавца {0}"
        sentences_template['pay_require'] = ["Обязательство Покупателя по уплате Цены будет исполнено надлежащим образом с момента списания денежных средств с корреспондентского счета банка Покупателя для их зачисления на счет Продавца."]
    elif 'нал' in pay_method:
        sentences_template['pay_phrase']  = "Уплата цены Договора производится путем передачи Покупателем наличных денежных средств Продавцу {0}"
        sentences_template['pay_require'] = ['Обязательство по уплате Цены считается исполненным в момент передачи денежных средств Продавцу.']
    else:
        sentences_template['pay_phrase']  = "Уплаты цены Договора производится путем передачи Покупателем имущества или прав требований Продавцу {0}"
        sentences_rephrase['pay_require'] = [f"Обязательство по уплате Цены считается исполненным в момент передачи  {get_my(json_result, 'Способ оплаты товара')} Продавцу."]
    
    fast_select('pay_phrase', 'Момент оплаты услуги', 
                sentences_template['pay_phrase'].format(f"в срок {get_my(json_result, 'Момент оплаты услуги')}."),
                sentences_template['pay_phrase'].format(' в момент получения Покупателем Товара.'))
    
    #seller_must = get_my(json_result, 'Обязанности продавца').replace(get_my(json_result, 'Товар'), '')
    
    fast_select('обязанности_продавец', 'Обязанности продавца',
               f"Обязательства Продавца включают в себя {get_my(json_result, 'Обязанности продавца')}",
               None)
    fast_select('права_продавец', 'Права продавца',
               f"Права Продавца включают в себя {get_my(json_result, 'Права продавца')}",
               None)
    fast_select('обязанности_покупатель', 'Обязанности покупателя',
               f"Обязательства Покупателя включают в себя {get_my(json_result, 'Обязанности покупателя')}",
               None)
    fast_select('права_покупатель', 'Права покупателя',
               f"Права Покупателя включают в себя {get_my(json_result, 'Права покупателя')}.",
               None)
    fast_select('doc_accept', "Документ подтверждения продажи",
                f"Покупатель, в соответствии с условиями настоящего Договора, имеет право требовать от Продавца передачи Товара в полном объеме, включая все виды и количество Товара, указанные в документах: {get_my(json_result, 'Документ подтверждения продажи')}, составляющих неотъемлемую часть настоящего Договора.",
                "Покупатель, в соответствии с условиями настоящего Договора, имеет право требовать от Продавца передачи Товара в полном объеме, включая все виды и количество Товара, указанные в документах, составляющих неотъемлемую часть настоящего Договора.")
    
    
    respon = f"В случае неисполнения или ненадлежащего исполнения Стороной своих обязательств по Договору, Сторона, допустившая нарушение, обязана: \n \
              (a) совершить все необходимые и возможные действия, направленные на уменьшение негативных последствий такого неисполнения; \n \
              (b) возместить пострадавшей Стороне все убытки, понесенные ею в результате указанного неисполнения, включая реальный ущерб и упущенную выгоду;"
    
    fast_select('ответственность', "Неустойка",
               respon + f"\n (c) уплатить неустойку в размере {get_my(json_result, 'Неустойка')}.",
               respon)
    
    ending_seller, ending_buyer, ending = get_my(json_result, 'Расторжение договора продавцом'), get_my(json_result, 'Расторжение договора покупателем'), ''
    
    if (ending_seller == ending_buyer) and (ending_seller not in ['', "Неизвестно"]):
        sentences_rephrase['измена'] = [f"Каждая из сторон имеет право расторгнуть Договор в одностороннем порядке, предварительно {get_my(json_result, 'Расторжение договора продавцом')}."]
    elif (ending_seller in ['', "Неизвестно"]) and (ending_buyer in ['', "Неизвестно"]):
        sentences_template['измена'] = ["Настоящий Договор может быть расторгнут по соглашению Сторон, а также по иным основаниям, предусмотренным действующим законодательством Российской Федерации."]
    else:
        sentences_rephrase['измена'] = ''
        if ending_seller not in ['', "Неизвестно"]:
            sentences_rephrase['измена'] += f"Продавец обязан уведомить Покупателя о расторжении Договора {get_my(json_result, 'Расторжение договора продавцом')}."
        if ending_buyer not in ['', "Неизвестно"]:
            sentences_rephrase['измена'] += f"Покупатель обязан уведомить Продавцом о расторжении Договора {get_my(json_result, 'Расторжение договора покупателем')}."
        sentences_rephrase['измена'] = [sentences_rephrase['измена']]
    
    #nest_asyncio.apply()
    loop       = asyncio.get_event_loop()
    sentences  = loop.run_until_complete(generate_concurrently(model, tokenizer, sentences_rephrase, gen_config))
    sentences.update(sentences_template)
    
    head = sentences['шапка'][0].replace('\n', ' ')
    add_text(doc, head)
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ПРЕДМЕТ ДОГОВОРА", align='center')
    add_list(doc, head_number, ["По Договору Продавец обязуется передать Товар в собственность Покупателю, а Покупатель обязуется принять этот Товар и уплатить за него Цену.",
                                *sentences['предмет договора'], *sentences['доставка'],
                                "Если обязательства Продавца по передаче Товара были исполнены не в надлежащем виде, то Покупатель вправе отказаться от его исполнения в целом или в части, соответствующей непредставленному исполнению. В случае принятия исполнения в части Договор считается действительным и исполнимым в соответствующей части."
                               ])
    #--------------------------------------------------------------------------------------------------------

    add_heading_numbered_dkp(doc, head_number, "ЦЕНА", align='center')
    add_list(doc, head_number, [*sentences['цена'], * sentences_template['pay_phrase'], *sentences_rephrase['pay_require']])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ПРАВА И ОБЯЗАННОСТИ ПРОДАВЦА", align='center')
    add_list(doc, head_number, ["Продавец имеет право требовать уплаты Цены за Товар в установленный Договором срок.",
                f"Продавец обязуется передать Покупателю Товар надлежащего качества и в обусловленном настоящим Договором ассортименте.",
                *sentences['обязанности_продавец'], *sentences['права_продавец']])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ПРАВА И ОБЯЗАННОСТИ ПОКУПАТЕЛЯ", align='center')
    add_list(doc, head_number, [f"Покупатель, в соответствии с условиями настоящего Договора, обязан принять Товар от Продавца в установленные сроки и оплатить его по установленной Договором Цене.",
              f"Покупатель имеет право на получение Товара в соответствии с ассортиментом и количеством, указанным в настоящем Договоре",
              *sentences['doc_accept'], f"Покупатель вправе требовать от Продавца передачи Товара в установленный настоящим Договором срок.",
              f"Покупатель вправе требовать от Продавца передачи Товара, свободного от любых прав третьих лиц и не являющегося предметом спора или судебного разбирательства.",
              *sentences['обязанности_покупатель'], *sentences['права_покупатель']])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ОТВЕТСТВЕННОСТЬ СТОРОН", align='center')
    add_list(doc, head_number, ["За неисполнение или ненадлежащее исполнение договорных обязательств Стороны несут имущественную ответственность в соответствии с правом Российской Федерации, если иное не предусмотрено Договором.",
                                *sentences['ответственность'], "Возмещение убытков не освобождает нарушившую Сторону от исполнения принятых на себя обязательств в натуре."])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "РАЗРЕШЕНИЕ СПОРОВ", align='center')
    add_list(doc, head_number, ["Отношения, вытекающие из настоящего Договора, и связанные с его заключением, исполнением, прекращением и расторжением подлежат регулированию в соответствии с правом РФ.",
                                "Стороны согласились, что все споры и разногласия Сторон, вытекающие из настоящего Договора, должны быть урегулированы в претензионном порядке, а также при помощи переговоров.",
              "При невозможности урегулирования споров и разногласий в претензионном порядке, такие споры и разногласия подлежат разрешению в судебном порядке.",
              "Любой спор, разногласие или претензия, вытекающие из Договора и возникающие в связи с ним, в том числе связанные с его нарушением, заключением, изменением, прекращением или недействительностью могут быть переданы на рассмотрение компетентного суда. Территориальная подсудность споров определяется по месту регистрации Ответчика."])
    #--------------------------------------------------------------------------------------------------------

    add_heading_numbered_dkp(doc, head_number, "КОНФИДЕНЦИАЛЬНОСТЬ", align='center')
    add_list(doc, head_number, ["В течение срока действия Договора и в течение 1 (одного) года после прекращения его действия каждая из Сторон обязуется не раскрывать третьим лицам без предварительного письменного согласия другой Стороны, а также не использовать в своих личных интересах или с любыми иными целями (за исключением связанных с исполнением Договора) какую-либо информацию (в письменной, устной или иной форме), которая была получена в связи с заключением или исполнением Договора («Конфиденциальная информация»).",
              f"Запрет, установленный пунктом 7.1 Договора, не распространяется на случаи, когда: \n \
              \t(a) информация была размещена в общедоступном источнике до даты Договора либо до того момента, когда она была предоставлена соответствующей Стороне, либо после того, как она была предоставлена соответствующей Стороне, однако в общедоступный источник такая информация попала иным образом, чем вследствие нарушения, допущенного соответствующей Стороной; \n \
              \t(b) информация не может относиться к конфиденциальной в силу положений применимого права; \n \
              \t(c) раскрытие информации является обязательным в целях соблюдения требований применимого права либо требований регулирующих органов; \n \
              \t(d) раскрытие должно быть осуществлено в целях совершения сделок или действий, предусмотренных положениями Договора или иного документа, ссылка на который содержится в Договоре; либо \n \
              \t(e) информация раскрывается Стороной своим работникам, аффилированным лицам и (или) консультантам при соблюдении последними конфиденциальности, как это предусмотрено настоящим Договором. ",
              f"В случае, предусмотренном пунктом 7.2(c), Сторона вправе раскрывать конфиденциальную информацию только в том объеме, который необходим и достаточен для соблюдения требований применимого права или требований регулирующих органов. ",
              "Стороны обязуются незамедлительно уведомлять друг друга обо всех фактах разглашения конфиденциальной информации третьим лицам, а также обо всех случаях, которые создают или могут создать угрозу разглашения конфиденциальной информации."])
    #--------------------------------------------------------------------------------------------------------
    
    
    add_heading_numbered_dkp(doc, head_number, "ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", align='center')
    add_list(doc, head_number, ["""Настоящий Договор вступает в силу с даты его подписания уполномоченными представителями обеих Сторон и действует до полного исполнения ими обязательств по настоящему Договору."""])
    add_list(doc, head_number, ["При расторжении Договора каждая Сторона обязана возвратить другой Стороне все полученное в рамках Договора.",
                       "Сторона не имеет право передавать или уступать любые права и обязательства по данному договору третьим лицам, за исключением случаев, когда это сделано после получения письменного разрешения второй стороны.",
                                "Любые изменения и дополнения к настоящему Договору должны быть составлены в письменной форме и подписаны Сторонами или надлежаще уполномоченными на то представителями Сторон.",
                                *sentences['измена'], "Договор составлен в двух экземплярах, имеющих равную юридическую силу, по одному для каждой из Сторон."])
    #--------------------------------------------------------------------------------------------------------
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Реквизиты продавца'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = 'Реквизиты покупателя'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    def dict_processing(data, row_cells, seller, term):
      if isinstance(data, dict):
        row_cells[seller].text = f"{term}: \n" + "\n".join([': '.join(item) for item in list(data.items())])
        run = row_cells[seller].paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
      else:
          row_cells[seller].text = term + ": " + str(data)
          run = row_cells[seller].paragraphs[0].runs[0]
          run.font.name = 'Times New Roman'

    for term in terms:
        row_cells = table.add_row().cells
        for i, side in enumerate(["продавца", "покупателя"]):
            cur_term = term
            if (json_result.get(f"Юридическая форма {side}") not in ["ООО",
                                    "Общество с ограниченной ответственностью", "АО", "Акционерное общество"]) and (term == 'ОГРН'): 
                cur_term = "ЕГРИП"
            print(json_result.get(f"Юридическая форма {side}"), cur_term)
            dict_processing(get_my(json_result, cur_term + f" {side}"), row_cells, i, cur_term)




    doc.add_paragraph('Подписи сторон:')

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Продавец'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    hdr_cells[1].text = 'Покупатель'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells = table.add_row().cells
    row_cells[0].text = f"_____/{get_my(json_result, 'ФИО продавца')}(подпись/Ф.И.О.)"
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells[1].text = f"_____/{get_my(json_result, 'ФИО покупателя')}(подпись/Ф.И.О.)"
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    
    doc.save(path_to_file + 'dcp.docx')




async def json_to_doc_uslugi(json_result, path_to_file, model, tokenizer, gen_config=None):
    global DEFAULT_SYSTEM_PROMPT_2 
    DEFAULT_SYSTEM_PROMPT_2 = """
    Ты - юридический эксперт. 
    Твоя задача - исправлять ошибки в юридическом тексте, 
    вернув в исправленном виде и не потеряв ключевые юридические сущности. 
    Не возвращай никакого другого текста, кроме отредактированного текста! 
    С большой буквы всегда: "Договор", "Цена", "Услуги", "Заказчик", "Исполнитель", "Сторона", "Расходы".
    Текст: 
    
    """
    
    doc = Document()
    today = date.today()
    data = today.strftime("%d.%m.%Y")
    head_number = create_multilevel_numbering(doc)
    add_heading_numbered_dkp(doc, None, 'ДОГОВОР ВОЗМЕЗДНОГО ОКАЗАНИЯ УСЛУГ')
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'г. Москва'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = '\t' + data
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
        
    sentences_template, sentences_rephrase = {}, {}
    def fast_select(key, entity, text_rep, text_temp):
        if get_my(json_result, entity) not in ['', "Неизвестно", 'неизвестно', 'нет данных']:
            sentences_rephrase[key] = [text_rep]
        else:
            sentences_template[key] = [text_temp]
    
    sentences_rephrase['шапка'] = [f"""{get_head_uslugi(json_result, "заказчика")} далее именуемый «Заказчик», с одной стороны, и \
                                       {get_head_uslugi(json_result, "исполнителя")} далее именуемый «Исполнитель», с другой стороны, заключили \
                                        настоящий договор (далее - «Договор») о нижеследующем:\n\n"""
                                  ]
    sentences_rephrase['предмет'] = [f"По заданию Заказчика Исполнитель обязуется {get_my(json_result, 'Наименование услуги')} (далее - «Услуги»)."]
    
    fast_select('адрес', 'Адрес оказания услуги', 
                f"Исполнитель оказывает услуги по адресу {get_my(json_result, 'Адрес оказания услуги')}'", 
                "Стороны согласовали, что Исполнитель оказывает Услуги в месте своего фактического нахождения, если иное не обусловлено характером и объемом оказываемых Заказчику Услуг.")
    
    fast_select('объем', 'Объем выполнения услуги', 
                f"По настоящему договору, Исполнитель обязуется оказать следующие Услуги: {get_my(json_result, 'Объем выполнения услуги')}",
                "Оказание услуг по настоящему Договору осуществляется Исполнителем на основании заданий Заказчика. Содержание задания или уточнения в задании излагаются Заказчиком в письменной форме и направляются в адрес Исполнителя по почте, посредством факсимильной или электронной связи по усмотрению Заказчика.")
     
    fast_select('качество', "Качество выполнения услуги",
               f"Качество Услуг должно соответствовать следующим требованиям: {get_my(json_result, 'Качество выполнения услуги')}",
               "Качество Услуг должно соответствовать требованиям, обычно предъяляемым к Услугам данного типа, а также обязательным требованиям, предусмотренным нормативными актами.")
    
    recover = ", а также возместить расходы, понесенные Исполнителем при оказании Услуг по настоящему Договору и (или) в связи с настоящим Договором (далее – «Расходы»)." if get_my(json_result, 'Возмещение убытков') else "."
    nds = f"Вознаграждение включает в себя НДС." if get_my(json_result, 'НДС') == 'да' else ''
    pay_freq    = get_my(json_result, 'Частота оплаты товара') if get_my(json_result, 'Частота оплаты товара') not in ['', "Неизвестно"] else 'фиксированной стоимостью'
    sentences_rephrase['цена']      = [f"За оказание услуг Заказчик обязуется заплатить вознаграждение в размере {get_my(json_result, 'Цена услуги')} {get_my(json_result, 'Валюта оплаты')} {pay_freq} за оказание Услуг (далее - «Цена»). " + nds] 
    sentences_template['награда']   = ["Заказчик обязан выплатить Исполнителю Вознаграждение в соответствии с условиями настоящего Договора" + recover]
    
    pay_when_phrase = "Заказчик обязуется оплатить оказанные услуги {0}. Исполнитель не имеет права на получение с Заказчика предусмотренных ст. 317.1 ГК РФ процентов за пользование суммой отсрочки (рассрочки) оплаты." 
    fast_select('цена_дата', 'Момент оплаты услуги', 
                pay_when_phrase.format(f"в срок {get_my(json_result, 'Момент оплаты услуги')}"),
                pay_when_phrase.format('в разумный срок с момента подписания Сторонами Акта сдачи приема услуг.' if get_my(json_result, 'Акт оказания услуг') == 'да' else "в разумный срок после фактического принятия им Услуг от Исполнителя."))
    
    pay_method  = get_my(json_result, 'Способ оплаты товара').replace('неизвестно', 'безналичным').replace('Неизвестно', 'безналичным')
    if 'безнал' in pay_method:
        sentences_template['pay_phrase']  = ["Уплата цены Договора производится путем безналичного перечисления денежных средств на счет Исполнителя."]
        sentences_template['pay_require'] = ["Обязательство Заказчика по уплате Цены будет исполнено надлежащим образом с момента списания денежных средств с корреспондентского счета банка Заказчика для их зачисления на счет Исполнителя."]
    elif 'нал' in pay_method:
        sentences_template['pay_phrase']  = ["Уплата цены Договора производится путем передачи Заказчиком наличных денежных средств Исполнителю."]
        sentences_template['pay_require'] = ['Обязательство по уплате Цены считается исполненным в момент передачи денежных средств Исполнителю.']
    else:
        sentences_template['pay_phrase']  = ["Уплаты цены Договора производится путем передачи Заказчиком имущества или прав требований Исполнителю."]
        sentences_rephrase['pay_require'] = [f"Обязательство по уплате Цены считается исполненным в момент передачи  {get_my(json_result, 'Способ оплаты труда')} Исполнителю."]
    
    fast_select('срок', 'Длительность исполнения заказа',
                f"Исполнитель обязуется оказать услуги, предусмотренные Договором, {get_my(json_result, 'Длительность исполнения заказа')}",
               "сполнитель обязуется оказать услуги, предусмотренные Договором, в течение разумного срока, с учетом характера объема Услуг оказываемых Заказчику.")
    
    fast_select('условия', 'Дополнительные условия оказания услуг',
                f"Стороны также согласовали следующие условия оказания Услуг: {get_my(json_result, 'Дополнительные условия оказания услуг')}",
                None)
    
    fast_select('undone', "Невыполнение заказа из-за Заказчика",
               f"В случае невозможности исполнения, возникшей по вине Заказчика, {get_my(json_result, 'Невыполнение заказа из-за Заказчика')}.",
               "В случае невозможности исполнения, возникшей по вине Заказчика, Услуги подлежат оплате в полном объеме, если иное не предусмотрено законом или договором возмездного оказания услуг.")
    
    if get_my(json_result, 'Акт оказания услуг') == 'да': 
        sentences_template['факт'] = [
         "По факту оказания Услуг Исполнитель направляет Заказчику 2 (два) экземпляра подписанного Исполнителем акта сдачи-приема оказанных услуг, по форме указанной в приложении №1 к Договору. (далее – «Акта сдачи-приема оказанных услуг»)",
         "Заказчик в течение трех рабочих дней подписывает оба экземпляра Акта сдачи-приема оказанных услуг и отправляет один экземпляр Исполнителем.",
         "Если Заказчик не согласен со сведениями, указанными в Акте сдачи-приема оказанных услуг, Заказчик обязан в течение 3 (трех) рабочих дней с момента получения Заказчиком Акта сдачи-приема оказанных услуг соответственно направить Исполнителю свои мотивированные возражения относительно сведений, указанных в Акте сдачи-приема оказанных услуг, изложенные в письменной форме.",
         "Если Исполнитель согласен с письменными возражениями Заказчика на Акт сдачи-приема оказанных услуг, он обязан устранить недостатки услуг за свой счет и повторно направить Заказчику Акт сдачи-приема оказанных услуг."
        ] 
    else:
        sentences_template['факт'] = [
         "Моментом окончания оказания Услуг Исполнителем считается фактическое завершение оказания Услуг или истечение срока оказания Услуг, указанного в настоящем Договоре."
         "После фактического оказания Услуг или по истечению срока их оказания, Заказчик вправе требовать от Исполнителя устранения выявленных в оказанных Услугах недостатков за счет Исполнителя в течение месяца."
        ]

    ending_seller, ending_buyer, ending = get_my(json_result, 'Расторжение договора исполнителем'), get_my(json_result, 'Расторжение договора заказчиком'), ''
    
    if (ending_seller == ending_buyer) and (ending_seller not in ['', "Неизвестно"]):
        sentences_rephrase['измена'] = [f"Каждая из сторон имеет право расторгнуть Договор в одностороннем порядке, предварительно {get_my(json_result, 'Расторжение договора заказчиком')}."]
    elif (ending_seller in ['', "Неизвестно"]) and (ending_buyer in ['', "Неизвестно"]):
        sentences_template['измена'] = ["Настоящий Договор может быть расторгнут по соглашению Сторон, а также по иным основаниям, предусмотренным действующим законодательством Российской Федерации."]
    else:
        sentences_rephrase['измена'] = ''
        if ending_seller not in ['', "Неизвестно"]:
            sentences_rephrase['измена'] += f"Исполнитель обязан уведомить Заказчика о расторжении Договора {get_my(json_result, 'Расторжение договора исполнителем')}."
        if ending_buyer not in ['', "Неизвестно"]:
            sentences_rephrase['измена'] += f"Заказчик обязан уведомить Исполнителя о расторжении Договора {get_my(json_result, 'Расторжение договора заказчиком')}."
        sentences_rephrase['измена'] = [sentences_rephrase['измена']]
    
    respon = f"В случае неисполнения или ненадлежащего исполнения Стороной своих обязательств по Договору, Сторона, допустившая нарушение, обязана: \n \
              (a) совершить все необходимые и возможные действия, направленные на уменьшение негативных последствий такого неисполнения; \n \
              (b) возместить пострадавшей Стороне все убытки, понесенные ею в результате указанного неисполнения, включая реальный ущерб и упущенную выгоду;"
    
    fast_select('ответственность', "Неустойка",
               respon + f"\n (c) уплатить неустойку в размере {get_my(json_result, 'Неустойка')}.",
               respon)
    
    loop       = asyncio.get_event_loop()
    sentences  = loop.run_until_complete(generate_concurrently(model, tokenizer, sentences_rephrase, gen_config))
    sentences.update(sentences_template)
    
    #--------------------------------------------------------------------------------------------------------
    
    head = sentences['шапка']
    add_text(doc, head)
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ПРЕДМЕТ ДОГОВОРА", align='center')
    add_list(doc, head_number, [*sentences['предмет'], *sentences['объем'], *sentences['адрес'], *sentences['качество']])
    #--------------------------------------------------------------------------------------------------------

    add_heading_numbered_dkp(doc, head_number, "Цена услуг и порядок оплаты", align='center')
    add_list(doc,head_number, [*sentences['цена'], *sentences['награда'], *sentences['pay_phrase'], *sentences['цена_дата']]) 
    #--------------------------------------------------------------------------------------------------------
    
    if get_my(json_result, 'Способ оказания услуги') in ['самостоятельно оказывать услугу', 'неизвестно']:
        dates = ["Исполнитель обязуется оказать Услуги лично. Исполнитель имеет право привлекать к работе третьих лиц (соисполнителей) только по согласованию с Заказчиком.",
                 "В случае если Заказчик будет привлекать к работе, выполняемой Исполнителем по настоящему Договору, других исполнителей, Исполнитель должен будет по поручению Заказчика согласовывать с ними возникающие в процессе работы вопросы, а также проекты договоров и иной документации, которые будут совместно выноситься на рассмотрение Заказчика."]
    else:
        dates = ["Для оказания Услуг Исполнитель имеет право привлекать к работе третьих лиц (соисполнителей) без согласования с Заказчиком. При этом Исполнитель несет полную материальную ответственность перед Заказчиком за действия привлекамемых им третьих лиц.",
                 "Исполнитель несет перед Заказчиком ответственность за последствия неисполнения или ненадлежащего исполнения обязательств субисполнителем в соответствии с правилами пункта 1 статьи 313 и статьи 403 ГК РФ. а также оплачивает все расходы по привлечению указанных лиц."]
    
    add_heading_numbered_dkp(doc, head_number, "Сроки и условия оказания услуг", align='center')
    add_list(doc,head_number, [*sentences['срок'], *sentences['условия'], *dates, *sentences['undone']])
  #--------------------------------------------------------------------------------------------------------    
   
    add_heading_numbered_dkp(doc, head_number, "ПОДТВЕРЖДЕНИЕ ФАКТА ОКАЗАНИЯ УСЛУГ", align='center')
    add_list(doc,head_number, sentences['факт'])
  # --------------------------------------------------------------------------------------------------------

    add_heading_numbered_dkp(doc, head_number, "ОТВЕТСТВЕННОСТЬ СТОРОН", align='center')
    add_list(doc, head_number, ["За неисполнение или ненадлежащее исполнение договорных обязательств Стороны несут имущественную ответственность в соответствии с правом Российской Федерации, если иное не предусмотрено Договором.",
             *sentences['ответственность'], "Возмещение убытков не освобождает нарушившую Сторону от исполнения принятых на себя обязательств в натуре."])
     # --------------------------------------------------------------------------------------------------------

    if get_my(json_result, 'Расторжение договора') == 'по решению одной стороны': 
        break_contract = ['Сторона, желающая расторгнуть Договор, должна письменно уведомить об этом другую Сторону. В этом случае Договор считается расторгнутым через 30 (тридцать) дней с момента получения уведомления другой Стороной.', 'Сторона, желающая расторгнуть договор, обязана возместить другой стороне убытки, причиненные ей в связи с односторонним расторжением договора']
    else:
        break_contract = ['Договор не может быть в любое время расторгнут одной из Сторон', 'Действие настоящего Договора прекращается после выполнения всех своих обязательств Сторонами.']
    
    add_heading_numbered_dkp(doc, head_number, "РАЗРЕШЕНИЕ СПОРОВ", align='center')
    add_list(doc, head_number, ["Отношения, вытекающие из настоящего Договора, и связанные с его заключением, исполнением, прекращением и расторжением подлежат регулированию в соответствии с правом РФ.",
                                "Стороны согласились, что все споры и разногласия Сторон, вытекающие из настоящего Договора, должны быть урегулированы в претензионном порядке, а также при помощи переговоров.",
              "При невозможности урегулирования споров и разногласий в претензионном порядке, такие споры и разногласия подлежат разрешению в судебном порядке.",
              "Любой спор, разногласие или претензия, вытекающие из Договора и возникающие в связи с ним, в том числе связанные с его нарушением, заключением, изменением, прекращением или недействительностью могут быть переданы на рассмотрение компетентного суда. Территориальная подсудность споров определяется по месту регистрации Ответчика."])
    #--------------------------------------------------------------------------------------------------------

    add_heading_numbered_dkp(doc, head_number, "КОНФИДЕНЦИАЛЬНОСТЬ", align='center')
    add_list(doc, head_number, ["В течение срока действия Договора и в течение 1 (одного) года после прекращения его действия каждая из Сторон обязуется не раскрывать третьим лицам без предварительного письменного согласия другой Стороны, а также не использовать в своих личных интересах или с любыми иными целями (за исключением связанных с исполнением Договора) какую-либо информацию (в письменной, устной или иной форме), которая была получена в связи с заключением или исполнением Договора («Конфиденциальная информация»).",
              f"Запрет, установленный пунктом 7.1 Договора, не распространяется на случаи, когда: \n \
              \t(a) информация была размещена в общедоступном источнике до даты Договора либо до того момента, когда она была предоставлена соответствующей Стороне, либо после того, как она была предоставлена соответствующей Стороне, однако в общедоступный источник такая информация попала иным образом, чем вследствие нарушения, допущенного соответствующей Стороной; \n \
              \t(b) информация не может относиться к конфиденциальной в силу положений применимого права; \n \
              \t(c) раскрытие информации является обязательным в целях соблюдения требований применимого права либо требований регулирующих органов; \n \
              \t(d) раскрытие должно быть осуществлено в целях совершения сделок или действий, предусмотренных положениями Договора или иного документа, ссылка на который содержится в Договоре; либо \n \
              \t(e) информация раскрывается Стороной своим работникам, аффилированным лицам и (или) консультантам при соблюдении последними конфиденциальности, как это предусмотрено настоящим Договором. ",
              f"В случае, предусмотренном пунктом 7.2(c), Сторона вправе раскрывать конфиденциальную информацию только в том объеме, который необходим и достаточен для соблюдения требований применимого права или требований регулирующих органов. ",
              "Стороны обязуются незамедлительно уведомлять друг друга обо всех фактах разглашения конфиденциальной информации третьим лицам, а также обо всех случаях, которые создают или могут создать угрозу разглашения конфиденциальной информации."])
    #--------------------------------------------------------------------------------------------------------
    
    
    add_heading_numbered_dkp(doc, head_number, "ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", align='center')
    add_list(doc, head_number, ["""Настоящий Договор вступает в силу с даты его подписания уполномоченными представителями обеих Сторон и действует до полного исполнения ими обязательств по настоящему Договору."""])
    add_list(doc, head_number, ["При расторжении Договора каждая Сторона обязана возвратить другой стороне всю полученную по договору сумму до момента его расторжения. Исполнитель обязан вернуть Заказчику Цену в течение 10 рабочих дней со дня расторжения Договора, указанного в отправленном Покупателю уведомлении, тогда как Заказчик должен вернуть продавцу товар обратно за свой счет в течение 2 рабочих дней после возврата ранее уплаченной Цены.",
                       "Сторона не имеет право передавать или уступать любые права и обязательства по данному договору третьим лицам, за исключением случаев, когда это сделано после получения письменного разрешения второй стороны.",
                                "Любые изменения и дополнения к настоящему Договору должны быть составлены в письменной форме и подписаны Сторонами или надлежаще уполномоченными на то представителями Сторон.",
                                *sentences['измена'], "При одностороннем расторжении Договора Сторона, желающая расторгнуть Договор, обязана возместить другой Стороне убытки, причиненные ей в связи с односторонним расторжением Договора.",
                                "Договор составлен в двух экземплярах, имеющих равную юридическую силу, по одному для каждой из Сторон."])
   #--------------------------------------------------------------------------------------------------------

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Реквизиты заказчика'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = 'Реквизиты исполнителя'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    def dict_processing(data, row_cells, seller, term):
        if isinstance(data, dict):
            row_cells[seller].text = f"{term}: \n" + "\n".join([': '.join(item) for item in list(data.items())])
            run = row_cells[seller].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
        else:
            row_cells[seller].text = term + ": " + str(data)
            run = row_cells[seller].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'

    for term in terms:
        row_cells = table.add_row().cells
        for i, side in enumerate(["заимодавца", "заемщика"]):
            print(json_result.get(f"Юридическая форма {side}"), term)
            cur_term = term
            if (json_result.get(f"Юридическая форма {side}") not in ["ООО",
                                    "Общество с ограниченной ответственностью", "АО", "Акционерное общество"]) and (term == 'ОГРН'): 
                cur_term = "ЕГРИП"
                print('into')
            print(side, cur_term)
            dict_processing(get_my(json_result, cur_term + f" {side}"), row_cells, i, cur_term)


    doc.add_paragraph('Подписи сторон:')

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Заказчик'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    hdr_cells[1].text = 'Исполнитель'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells = table.add_row().cells
    row_cells[0].text = f"_____/{get_my(json_result, 'ФИО заказчика')}(подпись/Ф.И.О.)"
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells[1].text = f"_____/{get_my(json_result, 'ФИО исполнителя')}(подпись/Ф.И.О.)"
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    doc.save(path_to_file + 'uslugi.docx')





#-------------------------------------------------
async def json_to_doc_zaym(json_result, path_to_file, model, tokenizer, gen_config=None):
    global DEFAULT_SYSTEM_PROMPT_2 
    DEFAULT_SYSTEM_PROMPT_2 = """
    Ты - юридический эксперт. 
    Твоя задача - исправлять ошибки в юридическом тексте, вернув в исправленном виде. 
    Не возвращай никакого другого текста, кроме отредактированного текста! 
    С большой буквы всегда: "Договор", "Займ", "Заимодавец", "Заемщик", "Сторона".
    Никаких комментариев, только отредактированный текст.
    Текст: 

    """
    doc = Document()
    today = date.today()
    data = today.strftime("%d.%m.%Y")
    head_number = create_multilevel_numbering(doc)
    add_heading_numbered_dkp(doc, None, 'ДОГОВОР ВОЗМЕЗДНОГО ОКАЗАНИЯ УСЛУГ')
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'г. Москва'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = '\t' + data
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    
    order = get_my(json_result, 'Порядок выплаты процентов').replace('неизвестно', 'процент').replace('', 'процент')
    if order == 'процент':
        order_sentence = 'В соответствии с условиями настоящего Договора, средства, возвращенные Заемщиком в счет погашения Займа, подлежат первоочередному зачислению на погашение процентов, начисленных на сумму Займа. Лишь после полного погашения процентов, оставшаяся часть возвращенных средств направляется на погашение основного долга по Займу.'
    elif order == 'тело займа':
        order_sentence = 'В соответствии с условиями настоящего Договора, средства, возвращенные Заемщиком в счет погашения Займа, подлежат первоочередному зачислению на погашение тела долга. Лишь после полного погашения тела долга, оставшаяся часть возвращенных средств направляется на погашение процентов по Займу.'
    else:
        order_sentence = None
    print(order_sentence)
    order_pay  = get_my(json_result, 'Частота выдачи займа').replace('неизвестно', 'одним').replace('', 'одним')
    way_to_pay = 'путем передачи денежных средств в пользу Заемщика' if get_my(json_result, 'Способ передачи займа') == 'наличными' else 'путем перечисления денежных средств на расчетный счет Заемщика указанный в настоящем Договоре ниже'
    order_pay_sentence = f'Предоставление Займа осуществляется Займодавцем {order_pay} платежом,'
    order_when_give    = f"Заимодавец обязуется предоставить всю сумму Займа в течение {get_my(json_result, 'Срок выдачи займа') if get_my(json_result, 'Срок выдачи займа').lower() not in ['', 'неизвестно'] else 'разумного срока'} с момента подписания договора." 
    order_pay_sentence += way_to_pay
    
    before_return = 'Заемщик имеет право досрочного погашения Займа в полном объеме или частично. Одновременно с досрочным погашением всей суммы Займа подлежат уплате проценты, начисленные за пользование суммой займа на дату досрочного погашения Займа включительно.' if get_my(json_result, 'Досрочный возврат') != 'False' else None
    print(get_my(json_result, 'Срок выдачи займа'))
    when_to_return = f"Возврат суммы займа производится Заёмщиком не позднее {get_my(json_result, 'Срок возврата займа')}." if get_my(json_result, 'Cрок возврата займа') != '' else 'В соответствии с условиями настоящего договора, средства, возвращенные Заемщиком в счет погашения Займа, подлежат первоочередному зачислению на погашение тела долга. Лишь после полного погашения тела долга, оставшаяся часть возвращенных средств направляется на погашение процентов по Займу.'   
    
    
    ending_seller, ending_buyer, ending = get_my(json_result, 'Расторжение договора заимодавцем'), get_my(json_result, 'Расторжение договора заемщиком'), []
    
    if (ending_seller == ending_buyer) and (ending_seller not in ['', "Неизвестно"]):
        ending = [f"Каждая из сторон имеет право расторгнуть Договор в одностороннем порядке, предварительно {get_my(json_result, 'Расторжение договора заимодавцем')}."]
    elif (ending_seller in ['', "Неизвестно"]) and (ending_buyer in ['', "Неизвестно"]):
        ending = ["Настоящий Договор может быть расторгнут по соглашению Сторон, а также по иным основаниям, предусмотренным действующим законодательством Российской Федерации."]
    else:
        if ending_seller not in ['', "Неизвестно"]:
            ending.append(f"Заимодавец обязан уведомить Заемщика о расторжении Договора {get_my(json_result, 'Расторжение договора заимодавцем')}.")
        if ending_buyer not in ['', "Неизвестно"]:
            ending.append(f"Заемщик обязан уведомить Заимодавца о расторжении Договора {get_my(json_result, 'Расторжение договора заемщиком')}.")
    
    # ежемесячно
    sentences = {
        'шапка' : [f"""{get_head_uslugi(json_result, "заимодавца")} («Займодавец»); c одной стороны, и {get_head_uslugi(json_result, "заемщика")}\
    («Заемщик»); совместно именуемые в дальнейшем «Стороны», а по отдельности – «Сторона», заключили настоящий договор о предоставлении процентного займа (далее – «Договор»):

"""],
        
        'предмет' : [
            f"Заимодавец на условиях настоящего Договора обязуется предоставить Заемщику денежный займ в размере {get_my(json_result, 'Сумма займа')}  {get_my(json_result, 'Валюта займа')} (далее – «Займ»), а Заемщик обязуется возвратить Заимодавцу Займ, а также уплатить Займодавцу проценты за пользование Займом в порядке и сроки, предусмотренные настоящим Договором.",
              f"За пользование Займом по настоящему Договору Заемщик уплачивает проценты по ставке: {get_my(json_result, 'Процент займа').replace('нет процента', '19%')}.",
                   #get_my(json_result, 'Частоты начисления процентов') if get_my(json_result, 'Частоты начисления процентов') != '' else 'При начислении процентов принимается фактическое количество дней в году, равное 365 (366) дням.',
                   # get_my(json_result, 'Дата начала начисления процентов') if get_my(json_result, 'Дата начала начисления процентов') != '' else "Днем предоставления Займа считается день поступления суммы Займа на Счет заемщика, днем исполнения Заемщиком своих обязательств по возврату суммы Займа (его части) и по уплате процентов за пользование Займом Заимодавцу считается день передачи денежных средств в пользу Заимодавца.",
                   f"Заем, предоставляемый Займодавцем Заемщику в рамках настоящего Договора, является {get_my(json_result, 'Цель займа').lower().replace('неизвестно', 'нецелевой')}."],
        
        'порядок' : [order_pay_sentence, order_sentence, order_when_give, when_to_return, before_return,
                f"Проценты за пользование суммой Займа выплачиваются Заемщиком {get_my(json_result, 'Частота возврата займа').lower().replace('', 'ежемесячно').replace('неизвестно', 'ежемесячно')}",
                    ],
        
        'неустойка' : [f"В случае нарушения Заемщиком срока возврата Займа (части Займа) и (или) срока уплаты процентов за пользование Займом, Заемщик уплачивает Займодавцу неустойку {get_my(json_result, 'Неустойка заемщика')}" if get_my(json_result, 'Неустойка заемщика') != '' else None,  
              f"В случае нарушения Заимодавцем срока выдачи Займа (части Займа), Заимодавец уплачивает Заемщику неустойку {get_my(json_result, 'Неустойка заимодавца')}" if get_my(json_result, 'Неустойка заимодавца') != '' else None],
        
        'заключение' : [*ending]
             
      
    }
    loop       = asyncio.get_event_loop()
    sentences  = loop.run_until_complete(generate_concurrently(model, tokenizer, sentences, gen_config))
    
    
    head = sentences['шапка']
    add_text(doc, head)
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "ПРЕДМЕТ ДОГОВОРА", align='center')
    add_list(doc, head_number,
             sentences['предмет'][:2] 
             + [f"Проценты за пользование суммой займа рассчитываются исходя из фактического количества дней пользования Займом. Отсчет периода начисления процентов начинается со дня, следующего за днем предоставления Займа, и заканчивается днем возврата Займа, при этом день возврата Займа включается в период начисления процентов."] 
             + sentences['предмет'][2:])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "Порядок предоставления и погашения Займа, уплаты процентов по Займу", align='center')
    add_list(doc, head_number, sentences['порядок'])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "Ответственность Сторон", align='center')
    add_list(doc, head_number,  
             ["В случае неисполнения или ненадлежащего исполнения Сторонами своих обязательств по настоящему Договору, они несут ответственность друг перед другом, установленную законодательством Российской Федерации и настоящим Договором."] + sentences['неустойка'])
    #--------------------------------------------------------------------------------------------------------
   
    add_heading_numbered_dkp(doc, head_number, "Форс-мажор", align='center')
    add_list(doc, head_number,  
             ["Обстоятельства непреодолимой силы – обстоятельства, возникшие после заключения настоящего Договора в результате событий чрезвычайного характера, которые Стороны не могли предвидеть или предотвратить.",
              f"При наступлении обстоятельств, указанных в п. 4.1 Договора, каждая Сторона должна без промедления известить о них в письменном виде другую Сторону с целью прийти к соглашению об изменениях условий данного Договора. Извещение должно содержать данные о характере обстоятельств и, по возможности, давать оценку их влияния на возможность исполнения Стороной своих обязательств по настоящему Договору. Направление такого извещения не освобождает Заемщика от обязанности исполнять условия данного Договора. В случае если наступившие обстоятельства непреодолимой силы не являются общеизвестными, факт наступления указанных обстоятельств должен быть подтвержден документом, выданным компетентным органом.",
              f"В случаях наступления обстоятельств, предусмотренных в п.4.1 Договора, срок выполнения Стороной обязательств по настоящему Договору отодвигается соразмерно времени, в течение которого действуют эти обстоятельства и их последствия."])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "Конфиденциальность", align='center')
    add_list(doc, head_number,  
             ["Условия настоящего Договора и соглашений (протоколов и т.п.) к нему (далее – «Конфиденциальная информация») конфиденциальны и не подлежат разглашению.","Стороны принимают все необходимые меры для того, чтобы их сотрудники, агенты, правопреемники без предварительного согласия другой Стороны не информировали третьих лиц о деталях данного Договора."])
    #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "Разрешение споров", align='center')
    add_list(doc, head_number,  
              ["Отношения, вытекающие из настоящего Договора, и связанные с его заключением, исполнением, прекращением и расторжением подлежат регулированию в соответствии с правом РФ.",
                                "Стороны согласились, что все споры и разногласия Сторон, вытекающие из настоящего Договора, должны быть урегулированы в претензионном порядке, а также при помощи переговоров.",
              "При невозможности урегулирования споров и разногласий в претензионном порядке, такие споры и разногласия подлежат разрешению в судебном порядке.",
              "Любой спор, разногласие или претензия, вытекающие из Договора и возникающие в связи с ним, в том числе связанные с его нарушением, заключением, изменением, прекращением или недействительностью могут быть переданы на рассмотрение компетентного суда. Территориальная подсудность споров определяется по месту регистрации Ответчика."])
     #--------------------------------------------------------------------------------------------------------
    
    add_heading_numbered_dkp(doc, head_number, "Заключительные положения", align='center')
    add_list(doc, head_number, ["""Настоящий Договор вступает в силу с даты его подписания уполномоченными представителями обеих Сторон и действует до полного исполнения ими обязательств по настоящему Договору."""])
    add_list(doc, head_number, ["При расторжении Договора каждая Сторона обязана возвратить другой Стороне все полученное в рамках Договора.",
                       "Сторона не имеет право передавать или уступать любые права и обязательства по данному договору третьим лицам, за исключением случаев, когда это сделано после получения письменного разрешения второй стороны.",
                                "Любые изменения и дополнения к настоящему Договору должны быть составлены в письменной форме и подписаны Сторонами или надлежаще уполномоченными на то представителями Сторон."] +
                                sentences['заключение'] + [
                                "Договор составлен в двух экземплярах, имеющих равную юридическую силу, по одному для каждой из Сторон."])
     #--------------------------------------------------------------------------------------------------------
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Реквизиты заимодавца'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'

    hdr_cells[1].text = 'Реквизиты заемщика'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    
    def dict_processing(data, row_cells, seller, term):
        if isinstance(data, dict):
            row_cells[seller].text = f"{term}: \n" + "\n".join([': '.join(item) for item in list(data.items())])
            run = row_cells[seller].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
        else:
            row_cells[seller].text = term + ": " + str(data)
            run = row_cells[seller].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'

    for term in terms:
        row_cells = table.add_row().cells
        for i, side in enumerate(["заимодавца", "заемщика"]):
            cur_term = term
            if (json_result.get(f"Юридическая форма {side}") not in ["ООО",
                                    "Общество с ограниченной ответственностью", "АО", "Акционерное общество"]) and (term == 'ОГРН'): 
                cur_term = "ЕГРИП"
            dict_processing(get_my(json_result, cur_term + f" {side}"), row_cells, i, cur_term)
    doc.add_paragraph('Подписи сторон:')

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Заимодавец'
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    hdr_cells[1].text = 'Заемщик'
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells = table.add_row().cells
    row_cells[0].text = f"_____/{get_my(json_result, 'ФИО заимодавца')}(подпись/Ф.И.О.)"
    run = hdr_cells[0].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    row_cells[1].text = f"_____/{get_my(json_result, 'ФИО заемщика')}(подпись/Ф.И.О.)"
    run = hdr_cells[1].paragraphs[0].runs[0]
    run.font.name = 'Times New Roman'
    doc.save(path_to_file + 'zaym.docx')


function_dict = {"dcp": json_to_doc_dcp,
                 "uslugi": json_to_doc_uslugi,
                "zaym": json_to_doc_zaym}
                 



def json_to_doc(json_result, doc_class, path_to_file, model, tokenizer, gen_config=None):
    if gen_config is not None:
        gen_config.temperature = 0.02
    return asyncio.run(function_dict[doc_class](json_result, path_to_file, model, tokenizer, gen_config))
